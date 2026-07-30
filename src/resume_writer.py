"""
Renders a tailored resume dict (see data/resume_base.json for shape) to a .docx
file, following the formatting conventions:
- Org/institution name bold, left-aligned; location right-aligned on the same line
- Role title italicized; dates right-aligned on the line below
- Single black/dark-grey color scheme, no color accents
- Bullet points with hanging indent
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH

DARK_GREY = RGBColor(0x33, 0x33, 0x33)
BLACK = RGBColor(0x00, 0x00, 0x00)
RIGHT_TAB_POSITION = Pt(468)  # ~6.5in usable width on a standard 1in-margin page


def _add_two_column_line(doc, left_text, right_text, bold_left=False, italic_left=False):
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB_POSITION, WD_TAB_ALIGNMENT.RIGHT)
    run_left = p.add_run(left_text)
    run_left.bold = bold_left
    run_left.italic = italic_left
    run_left.font.color.rgb = BLACK
    p.add_run("\t")
    run_right = p.add_run(right_text)
    run_right.font.color.rgb = DARK_GREY
    return p


def _add_bullets(doc, bullets):
    for bullet in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(bullet)
        run.font.size = Pt(10.5)
        run.font.color.rgb = BLACK


def _fix_zoom_setting(doc):
    """python-docx's default settings.xml omits the required zoom percent
    attribute, which some validators flag. Set it explicitly."""
    settings = doc.settings.element
    zoom = settings.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}zoom"
    )
    if zoom is not None:
        zoom.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}percent",
            "100",
        )


def _tighten_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    heading2 = doc.styles["Heading 2"]
    heading2.font.size = Pt(11.5)
    heading2.font.color.rgb = BLACK
    heading2.font.bold = True
    heading2.paragraph_format.space_before = Pt(8)
    heading2.paragraph_format.space_after = Pt(2)

    list_bullet = doc.styles["List Bullet"]
    list_bullet.font.size = Pt(9.5)
    list_bullet.paragraph_format.space_after = Pt(1)
    list_bullet.paragraph_format.space_before = Pt(0)


def write_resume_docx(resume: dict, output_path: str, job_title: str = ""):
    doc = Document()
    _fix_zoom_setting(doc)
    _tighten_styles(doc)

    for section in doc.sections:
        section.top_margin = Pt(32)
        section.bottom_margin = Pt(32)
        section.left_margin = Pt(46)
        section.right_margin = Pt(46)

    contact = resume["contact"]

    # Header
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(1)
    name_run = name_p.add_run(contact["name"])
    name_run.bold = True
    name_run.font.size = Pt(17)
    name_run.font.color.rgb = BLACK

    contact_line = " | ".join(
        v for v in [contact.get("email"), contact.get("phone"), contact.get("location"),
                    contact.get("linkedin"), contact.get("github")] if v
    )
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_p.add_run(contact_line)
    contact_run.font.size = Pt(9.5)
    contact_run.font.color.rgb = DARK_GREY

    if job_title:
        note_p = doc.add_paragraph()
        note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note_run = note_p.add_run(f"Tailored for: {job_title}")
        note_run.italic = True
        note_run.font.size = Pt(8.5)
        note_run.font.color.rgb = DARK_GREY

    # Summary (tailored per posting by tailor.py — this is prime ATS real estate)
    if resume.get("summary"):
        doc.add_heading("Summary", level=2)
        summary_p = doc.add_paragraph()
        summary_run = summary_p.add_run(resume["summary"])
        summary_run.font.size = Pt(10)
        summary_run.font.color.rgb = BLACK

    # Education
    doc.add_heading("Education", level=2)
    for edu in resume.get("education", []):
        _add_two_column_line(
            doc, edu["institution"], edu["location"], bold_left=True
        )
        _add_two_column_line(
            doc, edu["program"], edu["dates"], italic_left=True
        )
        if edu.get("notes"):
            _add_bullets(doc, edu["notes"])

    # Experience
    if resume.get("experience"):
        doc.add_heading("Experience", level=2)
        for exp in resume["experience"]:
            _add_two_column_line(
                doc, exp["organization"], exp["location"], bold_left=True
            )
            _add_two_column_line(
                doc, exp["role"], exp["dates"], italic_left=True
            )
            _add_bullets(doc, exp["bullets"])

    # Projects
    if resume.get("projects"):
        doc.add_heading("Projects", level=2)
        for proj in resume["projects"]:
            tech = ", ".join(proj.get("tech", []))
            _add_two_column_line(
                doc, proj["name"], proj.get("dates", ""), bold_left=True
            )
            _add_bullets(doc, proj["bullets"])
            if tech:
                tech_p = doc.add_paragraph()
                tech_run = tech_p.add_run(f"Tech: {tech}")
                tech_run.italic = True
                tech_run.font.size = Pt(9)
                tech_run.font.color.rgb = DARK_GREY

    # Leadership
    if resume.get("leadership"):
        doc.add_heading("Leadership", level=2)
        for role in resume["leadership"]:
            _add_two_column_line(doc, role["organization"], "", bold_left=True)
            _add_two_column_line(doc, role["role"], role["dates"], italic_left=True)
            _add_bullets(doc, role["bullets"])

    # Skills — merged into compact lines (rather than one paragraph per
    # category) so a resume with many skill categories still fits on one page.
    skills = resume.get("skills", {})
    if skills:
        doc.add_heading("Skills", level=2)
        for label, items in skills.items():
            if not items:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            label_run = p.add_run(f"{label.replace('_', ' ').title()}: ")
            label_run.bold = True
            skills_run = p.add_run(", ".join(items))
            skills_run.font.size = Pt(9.5)

    # Certifications
    if resume.get("certifications"):
        doc.add_heading("Certifications", level=2)
        p = doc.add_paragraph()
        p.add_run(" • ".join(resume["certifications"])).font.size = Pt(9.5)

    doc.save(output_path)
    return output_path
